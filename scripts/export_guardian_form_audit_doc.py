from pathlib import Path
from bs4 import BeautifulSoup
import markdown
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

project = Path('/home/ubuntu/connect-tutors-bd')
source = project / 'docs' / 'guardian-hire-tutor-form-audit.md'
output = project / 'docs' / 'guardian-hire-tutor-form-audit.docx'

html = markdown.markdown(source.read_text(encoding='utf-8'), extensions=['tables', 'fenced_code', 'sane_lists'])
soup = BeautifulSoup(html, 'html.parser')
doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)

styles = doc.styles
styles['Normal'].font.name = 'Arial'
styles['Normal'].font.size = Pt(10)
styles['Normal'].font.color.rgb = RGBColor(24, 61, 96)
for name, size, color in [('Title', 24, RGBColor(11, 109, 176)), ('Heading 1', 18, RGBColor(18, 110, 169)), ('Heading 2', 14, RGBColor(49, 91, 121)), ('Heading 3', 11, RGBColor(49, 91, 121))]:
    styles[name].font.name = 'Arial'
    styles[name].font.size = Pt(size)
    styles[name].font.bold = True
    styles[name].font.color.rgb = color


def add_inline(paragraph, node):
    if isinstance(node, str):
        paragraph.add_run(node)
        return
    if node.name in ('strong', 'b'):
        run = paragraph.add_run(node.get_text())
        run.bold = True
    elif node.name in ('em', 'i'):
        run = paragraph.add_run(node.get_text())
        run.italic = True
    elif node.name == 'code':
        run = paragraph.add_run(node.get_text())
        run.font.name = 'Courier New'
    elif node.name == 'a':
        run = paragraph.add_run(node.get_text())
        run.underline = True
    else:
        for child in node.children:
            add_inline(paragraph, child)


def add_table(node):
    rows = node.find_all('tr')
    if not rows:
        return
    cols = max(len(row.find_all(['th', 'td'])) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r, row in enumerate(rows):
        cells = row.find_all(['th', 'td'])
        for c, cell in enumerate(cells):
            target = table.cell(r, c)
            target.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            target.text = cell.get_text(' ', strip=True)
            for p in target.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(8.5)
            if cell.name == 'th' or r == 0:
                for p in target.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(23, 59, 96)
    doc.add_paragraph()


def walk(parent):
    for node in parent.children:
        if getattr(node, 'name', None) in ('h1', 'h2', 'h3'):
            p = doc.add_paragraph(style={'h1': 'Title', 'h2': 'Heading 1', 'h3': 'Heading 2'}[node.name])
            p.add_run(node.get_text(' ', strip=True))
        elif getattr(node, 'name', None) == 'p':
            p = doc.add_paragraph()
            for child in node.children:
                add_inline(p, child)
        elif getattr(node, 'name', None) == 'blockquote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(node.get_text(' ', strip=True))
            run.italic = True
            run.font.color.rgb = RGBColor(73, 105, 129)
        elif getattr(node, 'name', None) in ('ul', 'ol'):
            for li in node.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Bullet' if node.name == 'ul' else 'List Number')
                for child in li.children:
                    add_inline(p, child)
        elif getattr(node, 'name', None) == 'table':
            add_table(node)

walk(soup)
doc.core_properties.title = 'Guardian Hire a Tutor Form Audit'
doc.core_properties.subject = 'Section-wise form changes and recommendations'
doc.core_properties.author = 'Manus AI'
doc.save(output)
print(output)
